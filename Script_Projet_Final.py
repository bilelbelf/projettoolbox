import os
import tkinter as tk
from tkinter import messagebox, filedialog
import threading
import nmap
import paramiko
from numpy import loadtxt
import socket
import hashlib
import zxcvbn
from fpdf import FPDF
from docx import Document
import time
import vulners

class CVEcheck:
    def __init__(self) -> None:
        self.api_status = self._api_validator()

    def _api_validator(self):
        try:
            with open("vuln_api", "r") as file:
                self.api_key = file.readline().strip()
            try:
                self.api = vulners.VulnersApi(api_key=self.api_key)
                return True
            except Exception as e:
                print(f"API key {self.api_key} is not working. Please check it again. Error: {e}")
                return False
        except FileNotFoundError:
            print(f"File vuln_api not found. Please create it with your API key from https://vulners.com/")
            return False

    def vulnerability_check(self, service_info):
        try:
            res = self.api.find_exploit(service_info, limit=5)
            return res
        except Exception as e:
            print(f"An error occurred while searching for vulnerabilities: {e}")
            return []

    def print_vuln_res(self, results, ip, port, service_info):
        if len(results) == 0:
            return "No known exploit was found."

        output = f"\nVulnerability search results for {ip}:{port} -> {service_info}\n"
        for res in results:
            output += (
                f"IP Address : Port   |   {ip}:{port}\n"
                f"CVE Code(s)        |   {', '.join(res.get('cvelist', []))}\n"
                f"Title                    |   {res.get('title', '')}\n"
                f"Family                 |   {res.get('bulletinFamily', '')}\n"
                f"CVSS                   |   {res.get('cvss', {}).get('score', '')}\n"
                f"Link                      |   {res.get('href', '')}\n"
                "---------------------------------------------\n"
            )
        return output

def scan_vulns(ip, port_range, nmap_path):
    cve = CVEcheck()
    if not cve.api_status:
        return []

    os.environ["PATH"] += os.pathsep + nmap_path
    nm = nmap.PortScanner()
    nm.scan(ip, port_range, arguments='-sV')

    results = []

    for proto in nm[ip].all_protocols():
        lport = nm[ip][proto].keys()
        for port in lport:
            service_name = nm[ip][proto][port]['name']
            service_version = nm[ip][proto][port].get('version', 'unknown')
            service_info = f"{service_name} {service_version}"
            vulns = cve.vulnerability_check(service_info)
            results.append((ip, port, service_info, vulns))

    return results

class LoginWindow(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title('Authentification')
        self.geometry('400x250')
        self.configure(bg='#b0e0e6')

        self.username_label = tk.Label(self, text="Nom d'utilisateur :", bg='#b0e0e6', font=('Helvetica', 12))
        self.username_entry = tk.Entry(self, font=('Helvetica', 12))
        
        self.password_label = tk.Label(self, text="Mot de passe :", bg='#b0e0e6', font=('Helvetica', 12))
        self.password_entry = tk.Entry(self, show="*", font=('Helvetica', 12))
        
        self.login_button = tk.Button(self, text="Connexion", command=self.check_credentials, bg='#4682b4', fg='white', font=('Helvetica', 12, 'bold'))

        self.username_label.pack(pady=10)
        self.username_entry.pack(pady=5)
        self.password_label.pack(pady=10)
        self.password_entry.pack(pady=5)
        self.login_button.pack(pady=20)

    def check_credentials(self):
        entered_username = self.username_entry.get()
        entered_password = self.password_entry.get()

        if entered_username == "admin" and entered_password == "admin":
            messagebox.showinfo('Succès', 'Bienvenue admin!')
            self.withdraw()
            self.main_window = MainWindow(self)
            self.main_window.mainloop()
        else:
            messagebox.showerror('Erreur', 'Nom d\'utilisateur ou mot de passe incorrect.')

class MainWindow(tk.Toplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title('Toolbox de Sécurité')
        self.geometry('800x600')
        self.configure(bg='#b0e0e6')

        self.menu_label = tk.Label(self, text="MENU", bg='#b0e0e6', font=('Helvetica', 16, 'bold'))
        self.menu_label.pack(pady=20)

        self.center_frame = tk.Frame(self, bg='#b0e0e6')
        self.center_frame.pack(expand=True)

        self.create_button("Analyse de Mot de Passe", self.analyze_password_interface)
        self.create_button("Analyse de Port", self.analyze_port_interface)
        self.create_button("Test Bruteforce", self.bruteforce_interface)
        self.create_button("Scan de Vulnérabilité", self.vulnerability_scan_interface)

    def create_button(self, text, command):
        button = tk.Button(self.center_frame, text=text, command=command, bg='#4682b4', fg='white', font=('Helvetica', 12, 'bold'), width=30)
        button.pack(pady=10)

    def vulnerability_scan_interface(self):
        vuln_window = tk.Toplevel(self)
        vuln_window.title("Scan de Vulnérabilité")

        label_target = tk.Label(vuln_window, text="Rentrez votre IP pour le scan :", font=('Helvetica', 12))
        label_target.pack(pady=5)
        entry_target = tk.Entry(vuln_window, font=('Helvetica', 12))
        entry_target.pack(pady=5)
        
        label_port_range = tk.Label(vuln_window, text="Rentrez la plage de ports (ex: 1-1024) :", font=('Helvetica', 12))
        label_port_range.pack(pady=5)
        entry_port_range = tk.Entry(vuln_window, font=('Helvetica', 12))
        entry_port_range.pack(pady=5)

        text_results = tk.Text(vuln_window, height=15, width=80)
        text_results.pack(pady=5)

        def start_vuln_scan():
            target = entry_target.get()
            port_range = entry_port_range.get()
            if not target or not port_range:
                messagebox.showerror("Erreur", "Veuillez entrer une IP et une plage de ports.")
                return

            def scan_vulns_thread():
                results = scan_vulns(target, port_range, "C:\\Program Files (x86)\\Nmap")
                text_results.delete(1.0, tk.END)
                if not results:
                    text_results.insert(tk.END, "Aucun résultat. Vérifiez l'API ou l'IP.")
                    return
                for res in results:
                    text_results.insert(tk.END, f"{res[0]}:{res[1]} ({res[2]})\n")
                    for vuln in res[3]:
                        text_results.insert(tk.END, f"  - {vuln['title']} (CVSS: {vuln.get('cvss', {}).get('score', 'N/A')})\n")
                self.vuln_scan_results = results  # Save results for PDF generation

            threading.Thread(target=scan_vulns_thread).start()

        def generate_vuln_scan_report():
            if not hasattr(self, 'vuln_scan_results'):
                messagebox.showerror("Erreur", "Veuillez d'abord démarrer le scan.")
                return
            report_content = []
            for res in self.vuln_scan_results:
                for vuln in res[3]:
                    report_content.append(
                        (
                            f"IP Address : Port   |   {res[0]}:{res[1]}",
                            f"CVE Code(s)        |   {', '.join(vuln.get('cvelist', []))}",
                            f"Title                    |   {vuln.get('title', '')}",
                            f"Family                 |   {vuln.get('bulletinFamily', '')}",
                            f"CVSS                   |   {vuln.get('cvss', {}).get('score', '')}",
                            f"Link                      |   {vuln.get('href', '')}",
                            "---------------------------------------------"
                        )
                    )
            save_report(report_content, 'Scan de Vulnérabilité')

        btn_start_scan = tk.Button(vuln_window, text="Démarrer le Scan", command=start_vuln_scan, bg='#4682b4', fg='white', font=('Helvetica', 12, 'bold'))
        btn_start_scan.pack(pady=5)

        btn_generate_pdf = tk.Button(vuln_window, text="Générer le PDF", command=generate_vuln_scan_report, bg='#4682b4', fg='white', font=('Helvetica', 12, 'bold'))
        btn_generate_pdf.pack(pady=5)

    def analyze_port_interface(self):
        port_window = tk.Toplevel(self)
        port_window.title("Analyse de Port")

        label_target = tk.Label(port_window, text="Rentrez votre IP pour le scan :", font=('Helvetica', 12))
        label_target.pack(pady=5)
        entry_target = tk.Entry(port_window, font=('Helvetica', 12))
        entry_target.pack(pady=5)
        
        text_results = tk.Text(port_window, height=15, width=80)
        text_results.pack(pady=5)

        def start_port_scan():
            target = entry_target.get()
            if not target:
                messagebox.showerror("Erreur", "Veuillez entrer une IP.")
                return

            def scan_ports():
                os.environ["PATH"] += os.pathsep + "C:\\Program Files (x86)\\Nmap"
                nm = nmap.PortScanner()
                text_results.insert(tk.END, f"Scanning Target: {target}\n")
                text_results.insert(tk.END, "Starting...\n")
                nm.scan(target, '1-1024')
                results = []

                for host in nm.all_hosts():
                    for proto in nm[host].all_protocols():
                        lport = nm[host][proto].keys()
                        for port in lport:
                            state = nm[host][proto][port]['state']
                            service = nm[host][proto][port].get('name', 'unknown')
                            version = nm[host][proto][port].get('version', 'unknown')
                            text_results.insert(tk.END, f"Port {port}: {state} (Service: {service}, Version: {version})\n")
                            results.append((port, service, version, state))
                
                text_results.insert(tk.END, "Scan terminé.\n")
                self.port_scan_results = results  # Save results for PDF generation

            threading.Thread(target=scan_ports).start()

        def generate_port_scan_report():
            if not hasattr(self, 'port_scan_results'):
                messagebox.showerror("Erreur", "Veuillez d'abord démarrer le scan.")
                return
            report_content = [("Port", "Service", "Version", "État")]
            report_content.extend(self.port_scan_results)
            save_report(report_content, 'Analyse de Port')

        btn_start_scan = tk.Button(port_window, text="Démarrer le Scan", command=start_port_scan, bg='#4682b4', fg='white', font=('Helvetica', 12, 'bold'))
        btn_start_scan.pack(pady=5)

        btn_generate_pdf = tk.Button(port_window, text="Générer le PDF", command=generate_port_scan_report, bg='#4682b4', fg='white', font=('Helvetica', 12, 'bold'))
        btn_generate_pdf.pack(pady=5)

    def bruteforce_interface(self):
        brute_window = tk.Toplevel(self)
        brute_window.title("Brute Force SSH")

        label_ip = tk.Label(brute_window, text="Entrez l'IP de l'hôte :", font=('Helvetica', 12))
        label_ip.pack(pady=5)
        entry_ip = tk.Entry(brute_window, font=('Helvetica', 12))
        entry_ip.pack(pady=5)

        label_username = tk.Label(brute_window, text="Entrez le Nom d'utilisateur :", font=('Helvetica', 12))
        label_username.pack(pady=5)
        entry_username = tk.Entry(brute_window, font=('Helvetica', 12))
        entry_username.pack(pady=5)

        label_first_name = tk.Label(brute_window, text="Entrez le Prénom :", font=('Helvetica', 12))
        label_first_name.pack(pady=5)
        entry_first_name = tk.Entry(brute_window, font=('Helvetica', 12))
        entry_first_name.pack(pady=5)

        label_dob = tk.Label(brute_window, text="Entrez la Date de Naissance (DDMMYYYY) :", font=('Helvetica', 12))
        label_dob.pack(pady=5)
        entry_dob = tk.Entry(brute_window, font=('Helvetica', 12))
        entry_dob.pack(pady=5)

        text_results = tk.Text(brute_window, height=15, width=80)
        text_results.pack(pady=5)

        def start_bruteforce():
            ip = entry_ip.get()
            username = entry_username.get()
            first_name = entry_first_name.get()
            dob = entry_dob.get()
            if not ip or not username or not first_name or not dob:
                messagebox.showerror("Erreur", "Veuillez remplir tous les champs.")
                return

            if not self.validate_ip(ip):
                messagebox.showerror("Erreur", "Adresse IP invalide.")
                return

            text_results.delete(1.0, tk.END)
            text_results.insert(tk.END, "Démarrage du brute force...\n\n")

            def bruteforce_ssh(ip):
                ssh = paramiko.SSHClient()
                ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())

                common_pass_path = r'C:\Users\bilel\Downloads\TEST OFFICIEL PYTHON\probable-v2-top12000.txt'

                try:
                    common_pass = loadtxt(common_pass_path, dtype=str)
                except FileNotFoundError:
                    messagebox.showerror("Erreur", f"Fichier {common_pass_path} introuvable.")
                    return

                results = []
                possible_passwords = generate_password_combinations(username, first_name, dob)
                possible_passwords.extend(common_pass)

                for password in possible_passwords:
                    text_results.insert(tk.END, f"Essai - Nom d'utilisateur: {username}, Mot de passe: {password}\n")
                    text_results.see(tk.END)
                    try:
                        ssh.connect(ip, username=username, password=password, timeout=5)
                        text_results.insert(tk.END, f"Succès - Nom d'utilisateur: {username}, Mot de passe: {password}\n")
                        results.append((username, password, "Succès"))
                        ssh.close()
                        self.bruteforce_results = results  # Save results for PDF generation
                        return
                    except paramiko.AuthenticationException:
                        text_results.insert(tk.END, f"Échec - Nom d'utilisateur: {username}, Mot de passe: {password}\n")
                        results.append((username, password, "Échec"))
                    except paramiko.SSHException as e:
                        text_results.insert(tk.END, f"Erreur SSH: {str(e)}\n")
                        break
                    except socket.error as e:
                        text_results.insert(tk.END, f"Erreur de connexion: {str(e)}\n")
                        break
                    except Exception as e:
                        text_results.insert(tk.END, f"Erreur: {str(e)}\n")
                        break

                    brute_window.update_idletasks()
                    time.sleep(1)

                text_results.insert(tk.END, "Brute force terminé.\n")
                self.bruteforce_results = results  # Save results for PDF generation

            threading.Thread(target=bruteforce_ssh, args=(ip,)).start()

        def generate_bruteforce_report():
            if not hasattr(self, 'bruteforce_results'):
                messagebox.showerror("Erreur", "Veuillez d'abord démarrer le brute force.")
                return
            report_content = [("Nom d'utilisateur", "Mot de passe", "Résultat")]
            report_content.extend(self.bruteforce_results)
            save_report(report_content, 'Test Bruteforce')

        btn_start_bruteforce = tk.Button(brute_window, text="Démarrer Bruteforce", command=start_bruteforce, bg='#4682b4', fg='white', font=('Helvetica', 12, 'bold'))
        btn_start_bruteforce.pack(pady=5)

        btn_generate_pdf = tk.Button(brute_window, text="Générer le PDF", command=generate_bruteforce_report, bg='#4682b4', fg='white', font=('Helvetica', 12, 'bold'))
        btn_generate_pdf.pack(pady=5)

    def validate_ip(self, ip):
        try:
            socket.inet_aton(ip)
            return True
        except socket.error:
            return False

    def analyze_password_interface(self):
        analyze_window = tk.Toplevel(self)
        analyze_window.title("Analyse du Mot de Passe")

        frame = tk.Frame(analyze_window)
        frame.pack(padx=10, pady=10)

        label_password_analysis = tk.Label(frame, text="Mot de passe à analyser :", font=('Helvetica', 12))
        label_password_analysis.grid(row=3, column=0)

        global entry_password
        entry_password = tk.Entry(frame, font=('Helvetica', 12))
        entry_password.grid(row=3, column=1)

        btn_analyze = tk.Button(frame, text="Analyser Mot de Passe", command=self.analyze_password, bg='#4682b4', fg='white', font=('Helvetica', 12, 'bold'))
        btn_analyze.grid(row=4, columnspan=2, pady=10)

        global result_text
        result_text = tk.Text(frame, height=10, width=50)
        result_text.grid(row=5, columnspan=2, pady=10)

        def generate_report():
            analysis = self.password_analysis(entry_password.get())
            report_content = [("Aspect", "Valeur")]
            report_content.append(("Force du mot de passe", f"{analysis['score']}/4"))
            weaknesses = self.check_password_strength(analysis)
            if weaknesses:
                for weakness in weaknesses:
                    report_content.append(("Faiblesse", weakness))
            else:
                report_content.append(("Faiblesse", "Aucune"))
            if self.is_password_reused(entry_password.get()):
                report_content.append(("Réutilisation", "Oui"))
            else:
                report_content.append(("Réutilisation", "Non"))

            if analysis['score'] < 3:
                report_content.append(("Conseils", "Utilisez une combinaison de lettres majuscules et minuscules, de chiffres et de symboles. Évitez les mots de passe évidents et courants."))

            save_report(report_content, 'Analyse de Mot de Passe')

        btn_generate_report = tk.Button(frame, text="Générer le rapport", command=generate_report, bg='#4682b4', fg='white', font=('Helvetica', 12, 'bold'))
        btn_generate_report.grid(row=6, columnspan=2, pady=10)

    def analyze_password(self):
        password = entry_password.get()
        analysis = self.password_analysis(password)
        result_text.delete(1.0, tk.END)
        result_text.insert(tk.END, f"RESULTAT DE L'ANALYSE\n")
        result_text.insert(tk.END, f"Force du mot de passe : {analysis['score']}/4\n")

        weaknesses = self.check_password_strength(analysis)
        if weaknesses:
            result_text.insert(tk.END, "Faiblesses du mot de passe :\n")
            for weakness in weaknesses:
                result_text.insert(tk.END, weakness + "\n")
        else:
            result_text.insert(tk.END, "Le mot de passe est sécurisé.\n")

        if self.is_password_reused(password):
            result_text.insert(tk.END, "Ce mot de passe a déjà été utilisé.\n")
        else:
            result_text.insert(tk.END, "Ce mot de passe n'a pas été utilisé auparavant.\n")

        with open("resultat.txt", "a") as file:
            hashed_password = hashlib.sha256(password.encode()).hexdigest()
            file.write(hashed_password + "\n")

    def password_analysis(self, password):
        analysis = zxcvbn.zxcvbn(password)
        return analysis

    def check_password_strength(self, analysis):
        suggestions = analysis['feedback']['suggestions']
        weaknesses = []

        if analysis['score'] < 3:
            weaknesses.append("Le mot de passe est faible.")

        if "reuse" in suggestions:
            weaknesses.append("Le mot de passe est réutilisé.")

        if "compromised" in suggestions:
            weaknesses.append("Le mot de passe est potentiellement compromis.")

        return weaknesses

    def is_password_reused(self, password):
        try:
            with open("resultat.txt", "r") as file:
                hashed_password = hashlib.sha256(password.encode()).hexdigest()
                return hashed_password in file.read()
        except FileNotFoundError:
            return False

def generate_password_combinations(username, first_name, dob):
    day = dob[:2]
    month = dob[2:4]
    year = dob[4:]
    combinations = [
        first_name + day, first_name + month, first_name + year,
        first_name + day + month, first_name + month + year, first_name + day + year,
        first_name + day + month + "!", first_name + month + year + "@", first_name + day + year + "?"
    ]
    return combinations

def save_report(content, title):
    file_type = [("PDF files", "*.pdf"), ("DOCX files", "*.docx")]
    file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=file_type)
    if file_path:
        if file_path.endswith(".pdf"):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)

            pdf.set_font("Arial", size=16)
            pdf.cell(200, 10, txt=title, ln=True, align='C')
            pdf.ln(10)

            for row in content:
                for item in row:
                    pdf.multi_cell(0, 10, txt=str(item), border=0)
                pdf.ln(5)
            
            pdf.output(file_path)
        elif file_path.endswith(".docx"):
            doc = Document()
            doc.add_heading(title, 0)
            for row in content:
                for item in row:
                    doc.add_paragraph(str(item))
                doc.add_paragraph("---------------------------------------------")
            
            doc.save(file_path)
        messagebox.showinfo("Succès", "Rapport généré avec succès.")

if __name__ == '__main__':
    login_window = LoginWindow()
    login_window.mainloop()
